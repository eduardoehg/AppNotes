from tkinter import messagebox
from tkinter import filedialog
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def dimensionamento(janela, largura, altura, teto=0):
    largura_tela = janela.winfo_screenwidth()
    altura_tela = janela.winfo_screenheight()
    x = (largura_tela - largura) // 2
    y = (altura_tela - (altura + teto)) // 2
    return largura, altura, x, y


def limpar_tela(frame):
    for widget in frame.winfo_children():
        widget.destroy()


def msg_sucesso(msg):
    messagebox.showinfo(title='SUCESSO!', message=f'{msg}')


def msg_erro(msg):
    messagebox.showinfo(title='ERRO!', message=f'{msg}')


def gerar_pdf(titulo, nota, janela):
    janela.grab_set()
    diretorio = filedialog.askdirectory()
    nome_arquivo = f'{diretorio}/{titulo}.pdf'
    documento = SimpleDocTemplate(nome_arquivo, pagesize=letter)
    pdf = []

    # Título centralizado
    titulo_estilo = getSampleStyleSheet()["Title"]
    titulo_paragrafo = Paragraph(titulo, titulo_estilo)
    pdf.append(titulo_paragrafo)

    # Adicionar espaço
    pdf.append(Spacer(1, 20))

    # Separar o texto em parágrafos
    paragrafos = nota.split("\n")
    texto_estilo = getSampleStyleSheet()["Normal"]
    texto_estilo.alignment = 4
    texto_estilo.leftIndent = 10
    texto_estilo.rightIndent = 10

    for paragrafo in paragrafos:
        if paragrafo.strip():
            elemento = Paragraph(paragrafo, texto_estilo)
            pdf.append(elemento)
            pdf.append(Spacer(1, 12))

    if diretorio:
        try:
            documento.build(pdf)
            janela.grab_set()
            msg_sucesso('Documento Word Gerado com Sucesso!')
        except Exception as ex:
            janela.grab_set()
            msg_erro(f'Houve um erro ao gerar o Word\n\nMotivo:{ex}')


def gerar_word(titulo, nota, janela):
    diretorio = filedialog.askdirectory()
    nome_arquivo = f'{diretorio}/{titulo}.docx'
    documento = Document()

    # Título centralizado
    titulo_paragrafo = documento.add_paragraph()
    titulo_run = titulo_paragrafo.add_run(titulo)
    titulo_run.font.size = Pt(18)
    titulo_paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adicionar espaço
    documento.add_paragraph()

    # Adicionar parágrafos justificados
    paragrafos = nota.split("\n")
    for paragrafo in paragrafos:
        if paragrafo.strip():
            p = documento.add_paragraph(paragrafo)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    if diretorio:
        try:
            documento.save(nome_arquivo)
            janela.grab_set()
            msg_sucesso('Documento Word Gerado com Sucesso!')
        except Exception as ex:
            janela.grab_set()
            msg_erro(f'Houve um erro ao gerar o Word\n\nMotivo:{ex}')
